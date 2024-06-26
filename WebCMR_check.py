import pandas as pd
import logging
import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.common.exceptions import (NoSuchElementException,
                                        UnexpectedAlertPresentException,
                                        TimeoutException,
                                        ElementClickInterceptedException)
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from Completeness import Completeness
from typing import Union
from docx import Document

class WebCMR_check(Completeness):

    """
    This class is used to perform checks and gather HL7 error examples from the TST environment of 
    WebCMR.

    Purpose:
    The purpose of this class is to check the completeness values for each specified field in WebCMR 
    data. If a specific field has a completeness value less than the agreed upon threshold, the 
    program picks the first accession number of an instance where that field value is missing. 
    It then searches for the HL7 message associated with that accession number on the TST WebCMR 
    environment and populates a Word document with that HL7 message as an example of the different 
    types of errors.

    Algorithm:
    The algorithm of this class involves the following steps:

    1. Initialize the class by providing the necessary parameters.
    2. Create a Chrome webdriver object using Selenium to establish a connection to the TST WebCMR 
        environment.
    3. Navigate to the specified URL for WebCMR, such as TSTWebCMR or TRNWebCMR.
    4. Find the username and password elements on the login page and enter the provided login 
        credentials.
    5. Submit the login form and wait for the page to load.
    6. Navigate to the Incoming Message Monitor (IMM) page.
    7. Input the specified accession numbers into the search bar on the IMM page using Selenium.
    8. Click the search button to perform the search.
    9. Retrieve the HL7 messages associated with the accession numbers that have failed the 
        completeness checks.
    10. Create a Word document to store the HL7 error examples.
    11. Iterate through each accession number and its associated result test or distinguishing 
        criteria.
    12. If the distinguishing criteria is a threshold error, add the HL7 message as a threshold error 
        example to the Word document.
    13. If the distinguishing criteria is a date combination error, add the HL7 message as a date 
        combination error example to the Word document.
    14. Repeat steps 12 and 13 for each accession number and distinguishing criteria.
    15. Save the Word document with the HL7 error examples.

    Note: The algorithm assumes the presence of the `Completeness` class, which this class inherits 
    from, and that the necessary methods for completeness checks, logging, and HL7 extraction are 
    defined in the `Completeness` class.

    """

    def __init__(
        self, 
        username, 
        paswrd, 
        url = 'https://test-sdcounty.atlasph.com/TSTWebCMR/pages/login/login.aspx',
        *args,
        **kwargs
        ):
        super().__init__(*args, **kwargs)
        self.url = url 
        self.username = username
        self.paswrd = paswrd

    def login(self): 
        """
        Logs in the user to the website using the provided credentials and returns the webdriver object.

        Returns:
            webdriver: The webdriver object after successful login.

        Raises:
            NoSuchElementException: If the username or password elements cannot be found.
        """

        # create chrome webdriver object with the above options
        logging.info('Starting connection to webdriver')
        service : ChromeService = ChromeService(executable_path="chromedriver.exe")
        driver : webdriver = webdriver.Chrome(service=service)

        # go to TST website
        driver.get(self.url)

        # Find the username and password elements and enter login credentials
        # time.sleep(1)
        logging.info('Logging into TST')
        username : webdriver = driver.find_element(By.ID, value="txtUsername")
        username.send_keys(self.username)
        password : webdriver = driver.find_element(By.ID, value="txtPassword")
        password.send_keys(self.paswrd)
        # time.sleep(.5)
        password.send_keys(Keys.RETURN)
        
        return driver
    
    def acc_test_search(self, acc_num, driver,resultTest=None):
        """
        A function to search for an accession number in the Incoming Message Monitor.
        
        Parameters:
            acc_num (str): The accession number to search for.
            driver (WebDriver): The WebDriver object representing the browser session.
            resultTest (Optional): An optional parameter to specify a test result.
        
        Returns:
            WebDriver: The WebDriver object representing the browser session after the search is performed.
        """

        # navigate to IMM menu
        logging.info('Going to Incoming Message Monitor...')
        _ = self.nav2IMM(driver) 

        logging.info('Inputting accession numbers into search bar')
        acc_box : webdriver = self.multiFind(
            driver=driver,
            element_id= 'txtAccession',
            xpath='/html/body/form/div[3]/div/div/table[3]/tbody/tr[2]/td/table/tbody/tr[1]/td[8]/input'
        )
        acc_box.clear()
        acc_box.send_keys(str(acc_num))

        search_id : str = 'ibtnSearch'
        search_btn : webdriver = self.multiFind(
            driver=driver,
            element_id= search_id,
            xpath='/html/body/form/div[3]/div/div/table[3]/tbody/tr[2]/td/table/tbody/tr[4]/td/div/input[1]'
        )
        search_btn.click()
        return driver
    
    def get_hl7(self):
        """
        Method to go into IMM menu and conduct a search based on accession number and ResultTest.
        The program grabs example HL7 messages that have failed either date value checks or our examples
        of regions where they have 
        
        This function performs the following steps:
    
            1. Creates a word document object for HL7 reports.
            2. Adds a heading to the document.
            3. Calls the `combined_query_df` function to get information from both Demographics and 
                Lab and combines them into one DataFrame.
            4. Calls the `completeness_report` function to get the completeness reports for both 
                demographics and lab data.
            5. Finds exceptions with incorrect date combinations by calling the `date_check` function.
            6. Finds exceptions with less completeness than the allowed threshold by calling the 
                `threshold_search` function.
            7. Combines the results from step 5 and step 6 to get a list of accession numbers.
            8. Logs in to the TST environment to scrape HL7 messages that were flagged as missing or 
                incorrect info.
            9. Iterates over the list of accession numbers and performs the following actions for each accession number:
                - If the distinguifier is a string, puts threshold error HL7 examples in the word 
                document.
                - If the distinguifier is a list, puts date combination error HL7 examples in the 
                word document.
            10. Saves the word document with the name "HL7_Error.docx".
            11. Returns None.
        
        :param self: The current instance of the class.
        :return: None
            
        """
        
        # Creating word document object for hl7 reports
        doc : Document = Document()
        doc.add_heading('HL7 Error Examples')

        # calling combined query
        logging.info('Getting information from both Demographics and Lab and combining them into one DF')
        combined_query_df : pd.DataFrame = self.combined_query_df()
        demo_complete_df ,lab_complete_df = self.completeness_report() 

        # Get the list of Accession numbers from both date_check and threshold_search
        logging.info('Finding exceptions with incorrect date combinations')
        date_accession : list[tuple(str, int,Union[str, list])] = self.date_check(combined_query_df)
        logging.info('Finding exceptions with less completeness than allowed threshold')
        threshold_accession : list[tuple(str, int,Union[str, list])] = self.threshold_search(
            master_table=combined_query_df,
            demo_complete_df=demo_complete_df,
            lab_complete_df=lab_complete_df
        )
        accession_search : list = date_accession + threshold_accession

        # get driver: 
        driver : webdriver = self.login()
        logging.info('Scraping TST environment for HL7 messages that were flagged as missing or incorrect info...')
        for index, search_params in enumerate(accession_search):
            try:
            
                result_test : str = search_params[0]
                acc_num : int = search_params[1]
                distinguifier : Union[str, list] = search_params[2] # union is a method to type hint two types

                if isinstance(distinguifier, str):
                    logging.info(f'''
                    Putting threshold error HL7 examples in word doc
                    ACCESSION # : {acc_num}
                    '''
                                 )
                    self.hl7_extraction(
                        doc, 
                        accession_search, 
                        index, 
                        result_test, 
                        acc_num, 
                        heading='THRESHOLD ERROR', 
                        driver=driver
                    )  
                    #time.sleep(2)
                if isinstance(distinguifier, list):
                    # Need to tailor accession search variable to give a good header 
                    logging.info(f'''
                    Putting date combination error HL7 examples in word doc
                    ACCESSION # : {acc_num}
                    '''
                                 )
                    accession_search : list = [accession_search[0], accession_search[1], accession_search[2][1]]
                    self.hl7_extraction(
                        doc, 
                        accession_search,
                        index, 
                        result_test, 
                        acc_num, 
                        heading='DATE ERROR', 
                        driver=driver
                    )
                    #time.sleep(2)
            except UnexpectedAlertPresentException:
                continue
        logging.info('Putting all HL7 examples into docx ... ')
        doc.save("HL7_Error.docx")
        return

    def hl7_extraction(self, doc, accession_search, index, result_test, acc_num, heading , driver):
        """
        Extracts information from an HL7 document and adds it to a Word document.

        Parameters:
            doc (Word.Document): The Word document to which the extracted information will be added.
            accession_search (List[List[str]]): A list of lists containing information about the different accessions to search.
            index (int): The index of the current accession being searched.
            result_test (str): The result test to search for.
            acc_num (int): The accession number to search for.
            heading (str): The heading to add to the Word document.
            driver (webdriver): The webdriver instance used for accessing web elements.

        Returns:
            None
        """
        driver : webdriver = self.acc_test_search(
                    acc_num=acc_num, resultTest=result_test, driver=driver
                    )
        table : str = driver.find_element(By.ID, "divContentsArea").text       
        doc.add_heading(f'{heading}: {accession_search[index][2]}')
        doc.add_paragraph(table)
        pass 

    def multiFind(self, driver, element_id, xpath=None, field_name=None):
        """
        Finds and returns an element on a web page using the given driver and element ID.

        The function uses the WebDriverWait class from the selenium.webdriver.support.ui module
        to wait for a specified amount of time until the element is present on the page.
        It first tries to find the element using the ID locator strategy.
        If the element is not found within the specified timeout period, it waits for an additional
        1.5 seconds and then tries to find the element using the XPath locator strategy.
        If the element is still not found, it tries to find the element using the field name locator strategy.
        If the element is found, it is returned as a WebElement object.

        Args:
            driver (WebDriver): The web driver used to interact with the web page.
            element_id (str): The ID of the element to be found.
            xpath (str, optional): The XPath expression to locate the element (default: None).
            field_name (str, optional): The name of the field to locate the element (default: None).
            
        Returns:
            WebElement: The found element.
        """
    
        wait = WebDriverWait(driver, 2)
        try: 
            element_btn = wait.until(EC.presence_of_element_located((By.ID, element_id)))
        except TimeoutException:
            time.sleep(1.5)
            if xpath:
                try: 
                    element_btn = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
                except: 
                    pass
            if field_name:
                try: 
                    element_btn = wait.until(EC.presence_of_element_located((By.XPATH, field_name)))
                except: 
                    pass
            else: 
                logging.error(f"Cannot find {element_id} on web")
                raise NoSuchElementException(f"Cannot find {element_id} on web")
        except ElementClickInterceptedException: 
                self.alert_handling(driver)
        return element_btn
    
    def go_home(self, driver):
        """
        Clicks the home button and the investigators search button.

        Args:
            driver: The driver object used to interact with the web page.

        Returns:
            None
        """

        # Clicking home button 
        home_btn = self.multiFind(
            driver = driver, 
            element_id='FragTop1_lbtnHome',
            xpath= '/html/body/form/table[2]/tbody/tr/td[1]/div/a'
            )
        home_btn.click()

        # Clicking Search button (this is the investigators equivalent of my home button)
        investigators_search = self.multiFind(
            driver=driver,
            element_id='FragTop1_mnuMain-menuItem002',
            xpath='/html/body/form/table[2]/tbody/tr/td[2]/table[36]/tbody/tr/td[1]'
        )
        investigators_search.click()
        return
    
    def nav2IMM(self, driver):
        """
        Navigates to the IMM page in the website. Only works if user is on the homepage of website
        
        Args:
            driver: The WebDriver object used to interact with the browser.
            
        Returns:
            None
        """

        # Navigating back to home page
        self.go_home(driver)

        # Looking at Administrator dropdown menu
        wait = WebDriverWait(driver, 8)
        #dropdown_menu = wait.until(EC.presence_of_element_located((By.ID, "FragTop1_mnuMain-menuItem017")))
        dropdown_menu = self.multiFind(
            driver=driver,
            element_id= "FragTop1_mnuMain-menuItem017",
            xpath='/html/body/form/table[2]/tbody/tr/td[2]/table[36]/tbody/tr/td[6]'
        )
        dropdown_menu.click()

        # Hopefully clicking on incoming message monitor options
        # Wait for the second level dropdown to be present and then click on it
        #second_menu = "FragTop1_mnuMain-menuItem017-subMenu-menuItem009"
        second_level_dropdown = self.multiFind(
            driver, 
            element_id= "FragTop1_mnuMain-menuItem017-subMenu-menuItem009",
            xpath='/html/body/form/table[2]/tbody/tr/td[2]/table[16]/tbody/tr[9]/td'
        )
        second_level_dropdown.click()

        # Wait for the desired option to be present and then click on it
        imm = 'FragTop1_mnuMain-menuItem017-subMenu-menuItem009-subMenu-menuItem003'
        desired_option = self.multiFind(
            driver=driver,
            element_id=imm,
            xpath='/html/body/form/table[2]/tbody/tr/td[2]/table[5]/tbody/tr[3]/td'
        )
        desired_option.click()
        return
    
    def date_check(self, combined_query_df) -> list:
        """
        Check the dates in the given combined query dataframe for errors.

        Method to find all combinations of dates that violate logical lense.
        i.e) breaks this logic
                collection_date < received_date < result_date
        where '<' signifies earlier date

        Args:
            combined_query_df (DataFrame): The combined query dataframe containing the date columns.
        
        Returns:
            list: An array of tuples representing the date errors. Each tuple contains the following:
                - The result text
                - The accession number
                - An array with the error type and a detailed error message
        """

        # Array of accession for date errors
        date_errors = []
        # looking at every row and checking date conditions
        for _, row in combined_query_df.iterrows():
            # get timestamp variables
            spec_col_date = row['SPECCOLLECTEDDATE']
            spec_rec_date = row['SPECRECEIVEDDATE']
            result_date = row['RESULTDATE']

            # running checks on valid time combinations 
            if spec_col_date > spec_rec_date:
                date_errors.append(
                    (
                        row['RESULTTEXT'],
                        row['ACCESSIONNUMBER'],
                        ['SpecCollectDate Error (w/Recieve Date)',f'SpecCollectDate Error (w/Recieve Date) : {spec_col_date} > {spec_rec_date}']
                    )
                )
            if spec_rec_date > result_date:
                date_errors.append(
                    (
                        row['RESULTTEXT'],
                        row['ACCESSIONNUMBER'],
                        ['SpecRecieveDate Error (w/Result Date)',f'SpecRecieveDate Error (w/Result Date) : {spec_rec_date} > {result_date}']
                    )
                )
            if spec_col_date > result_date:
                date_errors.append(
                    (
                        row['RESULTTEXT'],
                        row['ACCESSIONNUMBER'], 
                        ['SpecCollectDate Error (w/Result Date)',f'SpecCollectDate Error (w/Result Date) : {spec_col_date} > {result_date}']
                    )
                )
            
        return date_errors
    
    def threshold_search(self, master_table, demo_complete_df ,lab_complete_df) -> list:
        """
        This method is meant to look at the completeness report of both lab and demographics data, 
        and compare it to a standard threshold that has been predetermined an hardcoded into the 
        program. If a field has a percent complete value less than the threshold, the program grabs 
        one specific combo of Accession Number and Result Test where that field has a blank entry 

        This function follows the following steps:
        1. Ensures that the indexes of the master table are not an issue in later analysis by 
        resetting them.
        2. Concatenates the demographic complete dataframe and the lab complete dataframe into one 
        dataframe called `combined_complete_df`.
        3. Resets the indexes of the `combined_complete_df`.
        4. Converts the 'Percent Complete' column of `combined_complete_df` to a float data type.
        5. Builds a dictionary `threshold_key_pair` that contains the threshold value for every 
        specific field of interest.
        6. Initializes a list `threshold_error` to store the threshold errors.
        7. Iterates over each row in the `combined_complete_df` dataframe.
        8. Retrieves the field of interest and the percent complete value for the current row.
        9. Assigns the threshold value for the field of interest in the `threshold_key_pair` dictionary.
        10. Checks if the percent complete value is less than the threshold value.
        11. If the condition is true, filters the master table to create a subset where the 
        field of interest has missing values (NaN or Null).
        12. Ensures that the length of the `master_subset` is greater than 0.
        13. Retrieves the first row from `master_subset` that contains NaN value for the field of 
        interest.
        14. Appends the result text, accession number, and field of interest to the `threshold_error` 
        list.
        15. Asserts that the length of `threshold_error` is less than 41 to ensure that only one 
        threshold error is recorded per field of interest.
        16. Returns the `threshold_error` list containing the threshold errors.

        Args:
            master_table (pd.DataFrame): The master table dataframe.
            demo_complete_df (pd.DataFrame): The demographic complete dataframe.
            lab_complete_df (pd.DataFrame): The lab complete dataframe.

        Returns:
            list: A list of threshold errors, each containing the result text, accession number, 
            and field of interest.
        """

        # making sure index's are not an issue in later analysis
        master_table.reset_index(inplace=True)
        
        
        # Going to concatenate into one df
        combined_complete_df : pd.DataFrame 
        combined_complete_df = pd.concat([demo_complete_df, lab_complete_df])
        combined_complete_df.reset_index(inplace=True)
        combined_complete_df['Percent Complete'] : pd.Series
        combined_complete_df['Percent Complete'] = combined_complete_df['Percent Complete'].astype(float)
        print(combined_complete_df)
        # building a dictionary that has the threshold value for every specific field of interest
        # At the same time checking if the threshold value is less than the percent complete
        # and then trying to find the accession numbers for one of the blank fields in the that category
        threshold_key_pair : dict = {}
        threshold_vals : list = [
            100,100,100,100,100,100,100,100,100,100,100,100,100,100,100,100,100,100,
            100,100,100,100,100,100,100,100,100,100,100,100,100,95,95,95,95,95,
            100,100,100,100
            ] # need marjorie to give me a list of thresholds similar to this 
        
        threshold_error : list = []
        for i, row in combined_complete_df.iterrows():
            col : str = row['Fields of Interest']
            pc : float = row['Percent Complete']
            threshold_key_pair[col] = threshold_vals[i]

            # if Percent complete is lower than threshold we need to look at a subset of master table
            # and pull out a accession number of a field where there is a nan value. We will look at 
            # only the first occurrence of nan value for this search 
            
            if pc < float(threshold_vals[i]):
                # filtering table on missing values using both NaN and Null for missing
                master_subset : pd.DataFrame = master_table[master_table[col].isna() | master_table[col].isnull()]

                assert len(master_subset) > 0
                
                # grabbing first row that contains NaN value of columns that 
                master_subset_row: pd.DataFrame = master_subset.iloc[0]

            # Appending the accession numbers that do not meet threshold criteria  
                threshold_error.append(
                    (
                    master_subset_row['RESULTTEXT'],
                    master_subset_row['ACCESSIONNUMBER'],
                    col
                    )
                )
        # there are 40 fields of interest, want to make sure we only grab 1 from each field that fails
        # Percent Complete threshold set by marjorie, So there couldn't be more than 40 elements total
        # in threshold_error list 
        assert len(threshold_error) < 41
        
        return threshold_error
    
